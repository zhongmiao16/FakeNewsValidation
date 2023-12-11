import docx
import pandas as pd

df = pd.read_excel("raw_stimuli.xlsx", sheet_name=1)

document = docx.Document()
for index, row in df.iterrows():
    print(index)
    txt = df.loc[index, "Grammar-Corrected Response"]
    if pd.isnull(txt):
        print(index)
        continue
    cond = df.loc[index, "Type"]
    if cond == "HUMAN":
        cond = f"{index} (C1)"
    else:
        cond = f"{index} (C2)"
    document.add_heading(cond, 0)
    document.add_paragraph(txt)
    document.add_page_break()

document.save("raw_stimuli.docx")
